from __future__ import annotations

import asyncio
import csv
import inspect
import time
import unittest
from datetime import UTC, date, datetime
from io import BytesIO, StringIO
from pathlib import Path
from unittest.mock import AsyncMock, patch
from zipfile import ZipFile

from docx import Document
from fastapi import FastAPI
from openpyxl import load_workbook

from app.api import reports
from app.reports.models import (
    DailySummaryAppointment,
    DailySummaryMetrics,
    DailySummaryReport,
)
from app.reports.renderers import (
    render_daily_summary_csv,
    render_daily_summary_docx,
    render_daily_summary_pdf,
    render_daily_summary_txt,
    render_daily_summary_xlsx,
)


OWNER_ID = "6a43c75072c43ebd5355030b"


def _report(
    *,
    locale: str = "en",
    client_name: str = "Client",
    service_name: str = "Service",
    notes: str = "Notes",
) -> DailySummaryReport:
    return DailySummaryReport(
        owner_id=OWNER_ID,
        report_date=date(2026, 8, 16),
        generated_at=datetime(2026, 8, 16, 12, 0, tzinfo=UTC),
        locale=locale,  # type: ignore[arg-type]
        metrics=DailySummaryMetrics(
            total_clients=10,
            total_services=4,
            total_appointments=20,
            appointments_on_date=1,
            scheduled_on_date=1,
            completed_on_date=0,
            cancelled_on_date=0,
        ),
        appointments=(
            DailySummaryAppointment(
                starts_at="2026-08-16T09:30:00+00:00",
                client_name=client_name,
                service_name=service_name,
                status="scheduled",
                notes=notes,
            ),
        ),
    )


class ReportRendererTests(unittest.TestCase):
    def test_pdf_renderer_returns_pdf_bytes_and_escapes_markup(self) -> None:
        payload = render_daily_summary_pdf(
            _report(
                client_name="<b>Client & Test</b>",
                notes="<script>alert(1)</script>",
            )
        )
        self.assertTrue(payload.startswith(b"%PDF"))
        self.assertGreater(len(payload), 500)

    def test_txt_renderer_is_utf8_and_localized(self) -> None:
        text = render_daily_summary_txt(
            _report(locale="hy")
        ).decode("utf-8")
        self.assertIn("Օրական ամփոփ հաշվետվություն", text)
        self.assertIn("Client", text)
        self.assertIn("2026-08-16", text)

    def test_csv_renderer_is_excel_friendly_and_formula_safe(self) -> None:
        payload = render_daily_summary_csv(
            _report(
                client_name='=HYPERLINK("x")',
                service_name="+SUM(1,1)",
                notes="@CMD",
            )
        )
        self.assertTrue(payload.startswith(b"\xef\xbb\xbf"))
        rows = list(csv.reader(StringIO(payload.decode("utf-8-sig"))))
        appointment = rows[-1]
        self.assertEqual(appointment[1], '\'=HYPERLINK("x")')
        self.assertEqual(appointment[2], "'+SUM(1,1)")
        self.assertEqual(appointment[4], "'@CMD")

        minus_payload = render_daily_summary_csv(
            _report(notes="-2+2")
        )
        minus_rows = list(
            csv.reader(
                StringIO(
                    minus_payload.decode("utf-8-sig")
                )
            )
        )
        self.assertEqual(minus_rows[-1][4], "'-2+2")

    def test_xlsx_renderer_creates_two_sheets_and_blocks_formulas(self) -> None:
        payload = render_daily_summary_xlsx(
            _report(client_name="=1+1", notes="-2+2")
        )
        workbook = load_workbook(BytesIO(payload), data_only=False)
        try:
            self.assertEqual(workbook.sheetnames, ["Overview", "Appointments"])
            appointments = workbook["Appointments"]
            self.assertEqual(appointments["B2"].value, "'=1+1")
            self.assertEqual(appointments["E2"].value, "'-2+2")
            self.assertNotEqual(appointments["B2"].data_type, "f")
        finally:
            workbook.close()

    def test_pdf_and_xlsx_renderers_are_deterministic(self) -> None:
        report = _report()

        pdf_first = render_daily_summary_pdf(report)
        xlsx_first = render_daily_summary_xlsx(report)

        time.sleep(1.1)

        pdf_second = render_daily_summary_pdf(report)
        xlsx_second = render_daily_summary_xlsx(report)

        self.assertEqual(pdf_first, pdf_second)

        def package_entries(payload: bytes) -> dict[str, bytes]:
            with ZipFile(BytesIO(payload)) as package:
                return {
                    name: package.read(name)
                    for name in sorted(package.namelist())
                }

        xlsx_entries_first = package_entries(xlsx_first)
        xlsx_entries_second = package_entries(xlsx_second)

        self.assertEqual(
            xlsx_entries_first,
            xlsx_entries_second,
        )
        self.assertIn(
            b"2026-08-16T12:00:00Z",
            xlsx_entries_first["docProps/core.xml"],
        )

    def test_docx_renderer_creates_valid_document_with_tables(self) -> None:
        payload = render_daily_summary_docx(_report(locale="fr"))
        document = Document(BytesIO(payload))
        all_text = "\n".join(p.text for p in document.paragraphs)
        self.assertIn("SalonFlow AI - Rapport quotidien", all_text)
        self.assertGreaterEqual(len(document.tables), 2)

    def test_renderer_modules_have_no_database_or_fastapi_dependency(self) -> None:
        root = Path("app/reports/renderers")
        forbidden = ("get_database", "db.", "pymongo", "motor.", "fastapi")
        for path in root.glob("*.py"):
            source = path.read_text(encoding="utf-8")
            for token in forbidden:
                self.assertNotIn(token, source, msg=f"{token} found in {path}")

    def test_openapi_exposes_all_five_formats_with_same_trusted_inputs(self) -> None:
        app = FastAPI()
        app.include_router(reports.router, prefix="/reports")
        schema = app.openapi()

        expected = {
            "pdf": "application/pdf",
            "txt": "text/plain",
            "csv": "text/csv",
            "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }

        for extension, media_type in expected.items():
            operation = schema["paths"][f"/reports/daily-summary/{extension}"]["get"]
            parameter_map = {
                item["name"]: item
                for item in operation["parameters"]
            }
            self.assertEqual(
                set(parameter_map),
                {"date", "locale", "authorization"},
            )
            self.assertNotIn("owner_id", parameter_map)
            self.assertEqual(parameter_map["authorization"]["in"], "header")
            self.assertIn(media_type, operation["responses"]["200"]["content"])
            self.assertIn("403", operation["responses"])
            self.assertIn("503", operation["responses"])

    def test_all_format_routes_keep_auth_and_entitlement_dependencies(self) -> None:
        functions = (
            reports.export_daily_summary_pdf,
            reports.export_daily_summary_txt,
            reports.export_daily_summary_csv,
            reports.export_daily_summary_xlsx,
            reports.export_daily_summary_docx,
        )
        for function in functions:
            signature = inspect.signature(function)
            self.assertIs(
                signature.parameters["auth"].default.dependency,
                reports.require_auth,
            )
            self.assertIs(
                signature.parameters["_entitlement"].default.dependency,
                reports.require_reports_entitlement,
            )

    def test_generic_export_builds_one_model_then_renders_requested_format(self) -> None:
        canonical = _report(locale="ru")
        builder = AsyncMock(return_value=canonical)
        database = object()

        async def run():
            with (
                patch.object(reports, "get_database", return_value=database),
                patch.object(reports, "build_daily_summary_report", builder),
            ):
                response = await reports.export_daily_summary_csv(
                    date_str="2026-08-16",
                    locale="ru",
                    auth={"admin_id": OWNER_ID},
                    _entitlement=None,
                )
                chunks = []
                async for chunk in response.body_iterator:
                    if isinstance(chunk, str):
                        chunk = chunk.encode()
                    chunks.append(chunk)
                return response, b"".join(chunks)

        response, payload = asyncio.run(run())

        builder.assert_awaited_once_with(
            database=database,
            owner_id=OWNER_ID,
            report_date=date(2026, 8, 16),
            locale="ru",
        )
        self.assertEqual(response.media_type, "text/csv")
        self.assertEqual(
            response.headers["content-disposition"],
            'attachment; filename="salonflow_daily_summary_ru_2026-08-16.csv"',
        )
        self.assertTrue(payload.startswith(b"\xef\xbb\xbf"))


if __name__ == "__main__":
    unittest.main()
