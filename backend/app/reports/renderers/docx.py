from __future__ import annotations

from io import BytesIO

from docx import Document

from app.reports.models import DailySummaryReport
from app.reports.renderers import (
    format_report_datetime,
    report_status,
    report_text,
)


def render_daily_summary_docx(report: DailySummaryReport) -> bytes:
    locale = report.locale
    metrics = report.metrics

    document = Document()
    document.add_heading(report_text(locale, "title"), level=0)
    document.add_paragraph(
        f'{report_text(locale, "report_date")}: {report.report_date.isoformat()}'
    )
    document.add_paragraph(
        f'{report_text(locale, "generated_at")}: {report.generated_at.strftime("%Y-%m-%d %H:%M UTC")}'
    )

    document.add_heading(report_text(locale, "overview"), level=1)

    overview = document.add_table(rows=1, cols=2)
    overview.style = "Table Grid"
    header = overview.rows[0].cells
    header[0].text = report_text(locale, "metric")
    header[1].text = report_text(locale, "value")

    for label, value in (
        (report_text(locale, "total_clients"), metrics.total_clients),
        (report_text(locale, "total_services"), metrics.total_services),
        (report_text(locale, "total_appointments"), metrics.total_appointments),
        (report_text(locale, "appointments_on_date"), metrics.appointments_on_date),
        (report_text(locale, "scheduled_on_date"), metrics.scheduled_on_date),
        (report_text(locale, "completed_on_date"), metrics.completed_on_date),
        (report_text(locale, "cancelled_on_date"), metrics.cancelled_on_date),
    ):
        cells = overview.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)

    document.add_heading(report_text(locale, "appointments"), level=1)

    if not report.appointments:
        document.add_paragraph(report_text(locale, "no_appointments"))
    else:
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        header = table.rows[0].cells
        header[0].text = report_text(locale, "start")
        header[1].text = report_text(locale, "client")
        header[2].text = report_text(locale, "service")
        header[3].text = report_text(locale, "status")
        header[4].text = report_text(locale, "notes")

        for item in report.appointments:
            cells = table.add_row().cells
            cells[0].text = format_report_datetime(item.starts_at)
            cells[1].text = item.client_name
            cells[2].text = item.service_name
            cells[3].text = report_status(locale, item.status)
            cells[4].text = item.notes

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()

# PHASE_63D_REPORT_DOCUMENT_RENDERER
def render_report_document_docx(document: object) -> bytes:
    from io import BytesIO

    from docx import Document

    from app.reports.contracts import ReportDocument

    if not isinstance(document, ReportDocument):
        raise TypeError("document must be a ReportDocument")

    def safe(value: object) -> str:
        text = "" if value is None else str(value)
        return f"'{text}" if text.startswith(("=", "+", "-", "@")) else text

    output = BytesIO()
    report = Document()
    report.core_properties.title = document.report_type
    report.core_properties.author = "SalonFlowAI"
    generated = document.generated_at.replace(tzinfo=None)
    report.core_properties.created = generated
    report.core_properties.modified = generated

    report.add_heading(document.report_type, level=0)
    report.add_paragraph(
        (
            f"{document.period.start_date.isoformat()} .. "
            f"{document.period.end_date.isoformat()} "
            f"({document.period.timezone})"
        )
    )
    metrics = report.add_table(rows=1, cols=2)
    metrics.rows[0].cells[0].text = "Metric"
    metrics.rows[0].cells[1].text = "Value"
    for key, value in document.metrics.items():
        cells = metrics.add_row().cells
        cells[0].text = safe(key)
        cells[1].text = safe(value)

    if document.columns:
        table = report.add_table(rows=1, cols=len(document.columns))
        for index, value in enumerate(document.columns):
            table.rows[0].cells[index].text = safe(value)
        for row in document.rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = safe(value)

    report.save(output)
    return output.getvalue()
